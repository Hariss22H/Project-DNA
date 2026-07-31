"""Default document extractors (Member 3 may replace via the service container)."""

from __future__ import annotations

import io
import logging
import re
from typing import Iterable

from app.core.exceptions import AppError
from app.services.ingestion.base import DocumentExtractor, DocumentType, ExtractedDocument

logger = logging.getLogger(__name__)


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


class TextDocumentExtractor(DocumentExtractor):
    def supports(self, file_type: DocumentType) -> bool:
        return file_type in {DocumentType.TXT, DocumentType.MARKDOWN}

    async def extract(
        self,
        *,
        file_name: str,
        file_type: DocumentType,
        content: bytes,
    ) -> ExtractedDocument:
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1", errors="replace")
        cleaned = clean_text(text)
        return ExtractedDocument(
            file_name=file_name,
            file_type=file_type,
            text=cleaned,
            metadata={"extractor": "text"},
            char_count=len(cleaned),
        )


class PdfDocumentExtractor(DocumentExtractor):
    def supports(self, file_type: DocumentType) -> bool:
        return file_type == DocumentType.PDF

    async def extract(
        self,
        *,
        file_name: str,
        file_type: DocumentType,
        content: bytes,
    ) -> ExtractedDocument:
        try:
            import fitz  # PyMuPDF
        except ImportError as exc:  # pragma: no cover
            raise AppError(
                "PDF extractor dependency missing",
                status_code=500,
                code="extractor_dependency_missing",
            ) from exc

        try:
            with fitz.open(stream=content, filetype="pdf") as doc:
                pages = [page.get_text("text") for page in doc]
                page_count = doc.page_count
        except Exception as exc:  # noqa: BLE001
            raise AppError(
                "Unable to parse PDF document",
                status_code=400,
                code="pdf_parse_error",
                details=str(exc),
            ) from exc

        cleaned = clean_text("\n\n".join(pages))
        return ExtractedDocument(
            file_name=file_name,
            file_type=file_type,
            text=cleaned,
            metadata={"extractor": "pymupdf"},
            page_count=page_count,
            char_count=len(cleaned),
        )


class DocxDocumentExtractor(DocumentExtractor):
    def supports(self, file_type: DocumentType) -> bool:
        return file_type == DocumentType.DOCX

    async def extract(
        self,
        *,
        file_name: str,
        file_type: DocumentType,
        content: bytes,
    ) -> ExtractedDocument:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover
            raise AppError(
                "DOCX extractor dependency missing",
                status_code=500,
                code="extractor_dependency_missing",
            ) from exc

        try:
            document = Document(io.BytesIO(content))
            paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        except Exception as exc:  # noqa: BLE001
            raise AppError(
                "Unable to parse DOCX document",
                status_code=400,
                code="docx_parse_error",
                details=str(exc),
            ) from exc

        cleaned = clean_text("\n".join(paragraphs))
        return ExtractedDocument(
            file_name=file_name,
            file_type=file_type,
            text=cleaned,
            metadata={"extractor": "python-docx"},
            char_count=len(cleaned),
        )


class CompositeDocumentExtractor(DocumentExtractor):
    """Routes extraction to the first registered extractor that supports the type."""

    def __init__(self, extractors: Iterable[DocumentExtractor] | None = None) -> None:
        self._extractors = list(
            extractors
            or (
                TextDocumentExtractor(),
                PdfDocumentExtractor(),
                DocxDocumentExtractor(),
            )
        )

    def supports(self, file_type: DocumentType) -> bool:
        return any(extractor.supports(file_type) for extractor in self._extractors)

    async def extract(
        self,
        *,
        file_name: str,
        file_type: DocumentType,
        content: bytes,
    ) -> ExtractedDocument:
        for extractor in self._extractors:
            if extractor.supports(file_type):
                result = await extractor.extract(
                    file_name=file_name,
                    file_type=file_type,
                    content=content,
                )
                if not result.text.strip():
                    raise AppError(
                        "Document contains no readable text",
                        status_code=400,
                        code="empty_document",
                    )
                return result

        raise AppError(
            f"Unsupported document type '{file_type.value}'",
            status_code=400,
            code="unsupported_file_type",
        )
