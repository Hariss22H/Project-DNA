import io
import docx
from fastapi import UploadFile, HTTPException
from typing import Dict, Any

class DOCXParserService:
    async def parse(self, file: UploadFile) -> Dict[str, Any]:
        """Extract text, tables, headings, and metadata from a DOCX file."""
        filename = file.filename or ""
        if not filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="File must be a DOCX")
            
        try:
            content = await file.read()
            doc = docx.Document(io.BytesIO(content))
            
            core_properties = doc.core_properties
            
            paragraphs = []
            headings = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                paragraphs.append(text)
                
                # Detect headings
                if para.style.name.startswith('Heading'):
                    headings.append({"level": para.style.name, "text": text})
            
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
                
            return {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "paragraphs": paragraphs,
                "headings": headings,
                "tables": tables_data,
                "images_metadata": [], # Advanced XML parsing needed for precise image metadata
                "document_properties": {
                    "created": str(core_properties.created) if core_properties.created else "",
                    "modified": str(core_properties.modified) if core_properties.modified else "",
                    "category": core_properties.category or "",
                    "comments": core_properties.comments or "",
                    "subject": core_properties.subject or "",
                    "keywords": core_properties.keywords or ""
                }
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error parsing DOCX: {str(e)}")

docx_parser_service = DOCXParserService()
